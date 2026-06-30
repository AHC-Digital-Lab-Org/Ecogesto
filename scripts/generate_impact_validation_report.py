from __future__ import annotations

import csv
import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
MATRIX_PATH = DOCS_DIR / "matriz-validacion-ecogestos.csv"
ECOGESTOS_PATH = ROOT / "material" / "ecogestos.json"
OUTPUT_DOCX = DOCS_DIR / "informe-validacion-calculos-ecogestos-ahc.docx"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_LEFT_RIGHT = 120

BRAND_GREEN = RGBColor(79, 148, 71)
BRAND_DARK = RGBColor(34, 42, 26)
BRAND_SAGE = RGBColor(83, 115, 88)
MUTED = RGBColor(90, 96, 94)
LIGHT_FILL = "F2F4F7"
SOFT_GREEN = "EAF3E8"
CALLOUT_FILL = "F6F8F5"
BORDER = "D7DEE5"
WHITE = "FFFFFF"

OFFICIAL_RE = re.compile(
    r"MITECO|IDAE|REE|FAO|GHG|ICAO|WRAP|AEMET|ACA|Ag[eè]ncia|Ministerio|BOE|Ley 7/2022",
    re.I,
)
TECHNICAL_RE = re.compile(
    r"Oxford|Poore|Nemecek|Ellen MacArthur|Energy Saving Trust|Ecoembes|OCU|AENOR|European Environment Agency|Ecolabel|NOVA|MSCI|Morningstar|Right to Repair|iFixit",
    re.I,
)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_MARGIN_TOP_BOTTOM, bottom=CELL_MARGIN_TOP_BOTTOM, left=CELL_MARGIN_LEFT_RIGHT, right=CELL_MARGIN_LEFT_RIGHT):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(0, tbl_grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def cant_split_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def paragraph_border_bottom(paragraph, color="4F9447", size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def paragraph_shading(paragraph, fill=CALLOUT_FILL):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def add_para(doc, text="", size=11, bold=False, italic=False, color=BRAND_DARK, before=0, after=6, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, color=BRAND_GREEN if level in (1, 2) else BRAND_SAGE)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.7, color=BRAND_DARK)
    return p


def add_number(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=10.7, color=BRAND_DARK)
    return p


def add_callout(doc, label, text, fill=CALLOUT_FILL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    paragraph_shading(p, fill)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.8, color=BRAND_GREEN, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.8, color=BRAND_DARK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    repeat_header(table.rows[0])
    cant_split_row(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=BRAND_DARK, bold=True)
    for row in rows:
        data_row = table.add_row()
        cant_split_row(data_row)
        cells = data_row.cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx in (0, 2) and len(str(value)) < 24:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=BRAND_DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def infer_midpoint(text):
    normalized = text.replace(".", "").replace(",", ".")
    numbers = [float(x) for x in re.findall(r"\d+(?:\.\d+)?", normalized)]
    if not numbers:
        return 0
    if len(numbers) == 1:
        return numbers[0]
    return (numbers[0] + numbers[1]) / 2


def classify_metric(text):
    lowered = text.lower()
    if "indirecto" in lowered or "incalculable" in lowered:
        return "Indirecto"
    has_co2 = bool(re.search(r"kg\s*co2|co2e|co2eq|tonelada|t\s*co2|co2/a", lowered))
    has_water = "litro" in lowered or "agua" in lowered
    has_plastic = "plástico" in lowered or "plastico" in lowered
    has_euros = "€" in text or "eur" in lowered or "euro" in lowered
    active = sum([has_co2, has_water, has_plastic, has_euros])
    if active > 1:
        return "Mixto"
    if has_co2:
        return "CO2e"
    if has_water:
        return "Agua"
    if has_plastic:
        return "Plástico"
    if has_euros:
        return "Euros"
    return "Pendiente"


def infer_confidence(source, impact):
    joined = f"{source} {impact}"
    if "pendiente" in joined.lower():
        return "Pendiente"
    if "indirecto" in impact.lower() or "incalculable" in impact.lower():
        return "Baja"
    if "estimación ahc" in source.lower() and not OFFICIAL_RE.search(source) and not TECHNICAL_RE.search(source):
        return "Media"
    if OFFICIAL_RE.search(source) or TECHNICAL_RE.search(source):
        return "Alta"
    if "indirecto" in impact.lower():
        return "Baja"
    return "Media"


def source_family(source):
    if OFFICIAL_RE.search(source):
        return "A - oficial/técnica"
    if TECHNICAL_RE.search(source):
        return "B - literatura/sector"
    if "estimación ahc" in source.lower():
        return "C - supuesto AHC"
    return "C - pendiente/mixta"


def recommended_action(metric, confidence, source):
    if metric == "Indirecto":
        return "No sumar al total; explicar como efecto indirecto."
    if metric in ("Agua", "Plástico"):
        return "Mantener unidad separada y validar factor específico."
    if confidence == "Alta" and OFFICIAL_RE.search(source):
        return "Cerrar fórmula visible con factor oficial actualizado."
    if confidence == "Alta":
        return "Validar rango con fuente primaria y dejar trazabilidad."
    if confidence == "Media":
        return "Mostrar como estimación con supuestos hasta revisión técnica."
    return "Dejar pendiente o cualitativo hasta tener fuente."


def load_data():
    with MATRIX_PATH.open("r", encoding="utf-8-sig", newline="") as f:
        matrix = list(csv.DictReader(f))
    with ECOGESTOS_PATH.open("r", encoding="utf-8") as f:
        ecogestos = json.load(f)
    ecogestos_by_code = {item["codigo"]: item for item in ecogestos}

    enriched = []
    for row in matrix:
        impact = row.get("impacto_estimado_original", "")
        source = row.get("fuente_original", "")
        metric = classify_metric(impact)
        confidence = infer_confidence(source, impact)
        enriched.append(
            {
                **row,
                "metric": metric,
                "confidence": confidence,
                "source_family": source_family(source),
                "midpoint": infer_midpoint(impact),
                "action": recommended_action(metric, confidence, source),
                "source_short": source[:95] + ("..." if len(source) > 95 else ""),
                "resumen": ecogestos_by_code.get(row["codigo"], {}).get("resumen", ""),
            }
        )
    return enriched


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BRAND_DARK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BRAND_GREEN, 16, 8),
        ("Heading 2", 13, BRAND_GREEN, 12, 6),
        ("Heading 3", 12, BRAND_SAGE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number", "List Bullet 2", "List Number 2"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.7)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("EcoGestos AHC")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header.add_run(" | Validación metodológica de cálculos")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Página ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(footer)


def add_cover(doc, total, categories, high_priority, deep_validation_count):
    add_para(doc, "INFORME METODOLÓGICO", size=11, bold=True, color=BRAND_GREEN, after=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    title = p.add_run("Validación de cálculos de impacto")
    set_run_font(title, size=24, color=BRAND_DARK, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    subtitle = p2.add_run("EcoGestos AHC - lógica científica, fuentes, límites y defensa del MVP")
    set_run_font(subtitle, size=13.5, color=BRAND_SAGE, bold=True)

    metadata_rows = [
        ("Cliente / proyecto", "Asociación Huella Cívica - EcoGestos AHC"),
        ("Documento", "Informe de validación metodológica para presentación"),
        ("Fecha", "18 de junio de 2026"),
        ("Alcance", f"{total} EcoGestos, {categories} categorías, matriz de cálculo y fuentes"),
        ("Estado", "MVP completo con estimaciones operativas auditables"),
    ]
    for label, value in metadata_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.7, color=BRAND_DARK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.7, color=BRAND_DARK)

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, "4F9447", "14")
    rule.paragraph_format.space_after = Pt(12)

    add_callout(
        doc,
        "Tesis defendible",
        "El MVP no debe presentarse como certificación externa de huella de carbono. Sí puede presentarse como un sistema de estimación transparente, trazable y auditable: cada resultado separa unidades, indica fuente, conserva la fórmula/factor usado y distingue confianza alta, media, baja o pendiente.",
        SOFT_GREEN,
    )

    table_rows = [
        ("EcoGestos analizados", str(total), "Inventario actual del MVP"),
        ("Categorías", str(categories), "Movilidad, energía, alimentación, consumo, residuos, agua, finanzas y comunidad"),
        ("Prioridad alta", str(high_priority), "Acciones donde conviene cerrar primero fórmula/factor"),
        ("Validación profunda inicial", str(deep_validation_count), "Top 10 recomendado en la metodología"),
    ]
    add_table(doc, ["Indicador", "Valor", "Lectura"], table_rows, [2600, 1300, 5460], 9.5, SOFT_GREEN)
    doc.add_page_break()


def add_executive_summary(doc):
    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_para(
        doc,
        "El objetivo de este documento es dejar preparada la explicación técnica de los cálculos del MVP de EcoGestos AHC. La idea fuerte no es prometer exactitud absoluta, porque eso requeriría una auditoría externa o ACV completos por producto, sino demostrar que el sistema calcula con una metodología razonable, documentada y revisable.",
    )
    add_para(
        doc,
        "La arquitectura actual funciona como una capa de estimación: toma un impacto base anual, clasifica la unidad afectada, aplica un porcentaje de ejecución y un multiplicador de frecuencia. Después comunica el resultado con nivel de confianza y tipo de impacto. Esta estructura es suficiente para un MVP educativo y de seguimiento siempre que se use el lenguaje adecuado.",
    )
    add_callout(
        doc,
        "Mensaje para presentación",
        "Los cálculos son estimaciones orientativas basadas en fuentes oficiales, literatura técnica y supuestos explícitos. No mezclamos CO2, agua, plástico y euros; los presentamos por separado para evitar conclusiones falsas.",
    )
    add_heading(doc, "Decisiones clave", 2)
    for item in [
        "Usar fuentes oficiales españolas cuando el gesto depende del contexto local: MITECO, REE, IDAE, ACA y normativa BOE.",
        "Usar fuentes internacionales reconocidas cuando el gesto pertenece a cadenas de valor, vuelos o alimentación: GHG Protocol, ICAO, FAO y WRAP.",
        "Marcar los supuestos AHC como estimación hasta revisión técnica, aunque sean útiles para priorizar y educar.",
        "Evitar sumar impactos indirectos financieros o comunitarios con CO2 directo salvo que exista una fórmula y un factor validado.",
    ]:
        add_bullet(doc, item)


def add_scope_and_architecture(doc):
    add_heading(doc, "2. Alcance y límites", 1)
    add_para(
        doc,
        "El alcance cubierto es el catálogo de 50 EcoGestos del MVP, la lógica de cálculo implementada, la matriz de validación y el discurso técnico para defenderlo. Quedan fuera la certificación externa, una auditoría ACV completa por EcoGesto, la conexión a mediciones reales de consumo y cualquier inscripción formal en registros oficiales de huella.",
    )
    add_table(
        doc,
        ["Dentro del MVP", "Fuera del MVP"],
        [
            ("Estimaciones de impacto por EcoGesto con fuente y confianza.", "Certificación externa de huella de carbono."),
            ("Separación de CO2e, agua, plástico y euros.", "ACV completo de cada producto o servicio sustituido."),
            ("Fórmula visible y snapshot del factor usado en resultados.", "Medición real automática de consumo de cada persona."),
            ("Matriz de prioridades para revisión profesional.", "Validación legal o científica emitida por tercero independiente."),
        ],
        [4680, 4680],
        9.3,
    )

    add_heading(doc, "3. Arquitectura de cálculo", 1)
    add_para(
        doc,
        "La lógica operativa implementada en el paquete compartido del MVP sigue una estructura simple y auditable. Primero extrae o recibe un impacto base anual. Después infiere la métrica principal del texto o campos estructurados: CO2e, agua, plástico, euros, mixto, indirecto o pendiente. Por último escala ese valor por el porcentaje de aplicación y la frecuencia declarada.",
    )
    add_callout(
        doc,
        "Fórmula visible",
        "impacto_anual = impacto_base_anual x porcentaje_aplicacion x multiplicador_frecuencia",
        "F8FBF7",
    )
    add_table(
        doc,
        ["Paso", "Qué hace", "Por qué es defendible"],
        [
            ("1. Valor base", "Extrae un número o rango del impacto estimado y usa el punto medio.", "Permite trabajar con rangos sin inventar precisión falsa."),
            ("2. Métrica", "Clasifica CO2e, agua, plástico, euros, mixto, indirecto o pendiente.", "Evita sumar unidades incompatibles."),
            ("3. Confianza", "Evalúa la fuente: oficial/técnica, literatura, AHC o indirecta.", "Comunica incertidumbre de forma visible."),
            ("4. Escalado", "Aplica porcentaje y frecuencia del plan personal.", "Adapta el cálculo a la ejecución real de la persona."),
            ("5. Snapshot", "Guarda factor y fórmula usados en el resultado.", "El histórico no cambia si la ficha se edita después."),
        ],
        [1150, 3900, 4310],
        9.1,
        SOFT_GREEN,
    )

    add_heading(doc, "Multiplicadores de frecuencia", 2)
    add_table(
        doc,
        ["Frecuencia", "Multiplicador", "Uso previsto"],
        [
            ("Única", "0,25", "Acciones puntuales con efecto anual parcial."),
            ("Diaria", "1,25", "Hábitos muy repetidos o con intensidad superior a semanal."),
            ("Semanal", "1,00", "Base estándar de comparación del MVP."),
            ("Mensual", "0,50", "Acciones recurrentes pero menos intensas."),
        ],
        [2500, 1800, 5060],
        9.4,
    )


def add_sources_method(doc):
    add_heading(doc, "4. Jerarquía de fuentes", 1)
    add_para(
        doc,
        "La defensa científica del MVP no depende de una sola fuente, sino de una jerarquía. En España, los factores de electricidad, movilidad y energía deben apoyarse primero en organismos oficiales. Para categorías con alta variabilidad, como alimentación o cadena de valor, conviene usar literatura y guías internacionales y mantener rangos.",
    )
    add_table(
        doc,
        ["Nivel", "Fuente tipo", "Uso en EcoGestos", "Cómo se comunica"],
        [
            ("A", "Organismos oficiales o metodologías reconocidas: MITECO, REE, IDAE, ACA, BOE, ICAO, GHG Protocol.", "Movilidad, energía, electricidad, agua, residuos, vuelos y marco de alcance 3.", "Cálculo validable o estimación con fuente oficial."),
            ("B", "Literatura científica y entidades sectoriales reconocidas: FAO, WRAP, Oxford/Poore & Nemecek, Ellen MacArthur, Ecoembes, OCU.", "Alimentación, desperdicio, moda, envases, consumo y reparación.", "Estimación técnica con rango y supuestos."),
            ("C", "Supuestos AHC, fuentes divulgativas o impactos indirectos.", "Hábitos, organización, comunidad y finanzas sostenibles.", "Estimación orientativa o impacto cualitativo."),
        ],
        [900, 3300, 2920, 2240],
        8.9,
        SOFT_GREEN,
    )

    add_heading(doc, "Regla de oro de agregación", 2)
    for item in [
        "Solo se suma lo que comparte unidad, periodo y calidad metodológica suficiente.",
        "kg CO2e/año se puede sumar con kg CO2e/año; litros de agua, kg de plástico y euros se agregan por separado.",
        "Los impactos indirectos se muestran como señales educativas o cualitativas, no como CO2 directo del plan.",
        "Cuando una ficha mezcla varios efectos, el informe debe mostrar desglose y no convertir todo a una unidad común sin factor explícito.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Niveles de confianza", 2)
    add_table(
        doc,
        ["Nivel", "Criterio", "Frase segura para enseñar"],
        [
            ("Alta", "Fuente oficial o técnica reconocida, fórmula simple y revisable.", "Cálculo validable con fuente identificada."),
            ("Media", "Fuente razonable o supuesto AHC con base técnica pendiente de cierre.", "Estimación con supuestos documentados."),
            ("Baja", "Impacto indirecto o variabilidad demasiado alta para sumar.", "Impacto orientativo no agregable."),
            ("Pendiente", "Falta factor, fórmula o fuente suficiente.", "Impacto pendiente de cálculo."),
        ],
        [1200, 4300, 3860],
        9.2,
    )


def add_category_methodology(doc):
    add_heading(doc, "5. Lógica por categoría", 1)
    category_blocks = [
        (
            "Movilidad y transporte",
            "La fórmula base es distancia evitada o compartida multiplicada por un factor de emisión del vehículo sustituido. Para vuelos se debe usar calculadora/metodología específica, porque la distancia, cabina y operación del vuelo cambian mucho el resultado.",
            "km evitados x factor kg CO2e/km x frecuencia anual",
            "MITECO, IDAE, ICAO y GHG Protocol.",
        ),
        (
            "Energía y vivienda",
            "Los gestos de electricidad o climatización se calculan con kWh evitados por factor de emisión eléctrico o energético. Si el gesto es de comportamiento, se usa rango y se documenta el supuesto.",
            "kWh evitados x factor kg CO2e/kWh",
            "REE, MITECO e IDAE.",
        ),
        (
            "Alimentación",
            "Debe tratarse con rangos. Dieta, desperdicio alimentario, carne de vacuno y opciones vegetales dependen de producto, origen y hábito previo. Aquí la precisión falsa sería un riesgo.",
            "ración o kg evitado x factor comparativo x frecuencia anual",
            "FAO, WRAP y literatura científica.",
        ),
        (
            "Consumo y compras",
            "Suele ser categoría de sustitución: compra evitada, reparación, segunda mano o préstamo. La variabilidad de producto obliga a trabajar con rangos o señales cualitativas.",
            "unidades evitadas x factor de producto o rango de referencia",
            "Literatura sectorial, economía circular y supuestos documentados.",
        ),
        (
            "Residuos y plásticos",
            "Conviene separar kg de plástico evitado, CO2e de gestión de residuos y evidencia de prevención. La normativa española respalda la prevención y reducción, pero no da por sí sola un factor universal.",
            "kg plástico evitado o kg residuo separado x factor específico",
            "BOE Ley 7/2022, MITECO y entidades sectoriales.",
        ),
        (
            "Agua",
            "El agua se muestra como litros/año. Solo debe convertirse a CO2e si se conoce energía asociada al calentamiento, bombeo o tratamiento, y se declara el supuesto.",
            "litros evitados por uso x usos anuales",
            "ACA, MITECO, AEMET y referencias de consumo doméstico.",
        ),
        (
            "Finanzas sostenibles y comunidad",
            "Son efectos indirectos. Pueden ser muy potentes narrativamente, pero no deben mezclarse con la huella directa individual si no hay metodología financiera específica y auditable.",
            "indicador indirecto o cualitativo; no agregable por defecto",
            "GHG Protocol como marco; fuentes financieras bajo revisión.",
        ),
    ]
    rows = [(name, logic, formula, sources) for name, logic, formula, sources in category_blocks]
    add_table(doc, ["Categoría", "Lógica", "Fórmula orientativa", "Fuentes"], rows, [1600, 3800, 2200, 1760], 8.4, SOFT_GREEN)


def add_validation_summary(doc, rows):
    add_heading(doc, "6. Estado de la matriz", 1)
    total = len(rows)
    category_counts = Counter(row["categoria"] for row in rows)
    metric_counts = Counter(row["metric"] for row in rows)
    confidence_counts = Counter(row["confidence"] for row in rows)
    source_family_counts = Counter(row["source_family"] for row in rows)
    high_priority = sum(1 for row in rows if row["prioridad_validacion"] == "Alta")

    add_para(
        doc,
        f"La matriz actual contiene {total} EcoGestos. Todos mantienen estado editorial de revisión o pendiente de cálculo final, lo cual es correcto para no sobreprometer. Aun así, la app ya puede usar una confianza inferida por fuente para enseñar el MVP: alta cuando hay fuente oficial/técnica, media cuando hay supuesto razonable y baja cuando el impacto es indirecto.",
    )
    add_table(
        doc,
        ["Lectura", "Resultado"],
        [
            ("EcoGestos totales", str(total)),
            ("Categorías cubiertas", str(len(category_counts))),
            ("EcoGestos con prioridad alta de validación", str(high_priority)),
            ("Métricas principales", ", ".join(f"{k}: {v}" for k, v in metric_counts.most_common())),
            ("Confianza inferida", ", ".join(f"{k}: {v}" for k, v in confidence_counts.most_common())),
            ("Familia de fuentes", ", ".join(f"{k}: {v}" for k, v in source_family_counts.most_common())),
        ],
        [3300, 6060],
        9.3,
        SOFT_GREEN,
    )

    add_heading(doc, "Top 10 para validación profunda", 2)
    add_para(doc, "Estos son los gestos que conviene cerrar primero porque combinan impacto, claridad de medición y valor pedagógico.")
    top_codes = [
        "ECO-MOV-002",
        "ECO-MOV-003",
        "ECO-MOV-004",
        "ECO-MOV-006",
        "ECO-ENE-004",
        "ECO-ENE-003",
        "ECO-ALI-001",
        "ECO-ALI-002",
        "ECO-RES-002",
        "ECO-AGU-002",
    ]
    by_code = {row["codigo"]: row for row in rows}
    top_rows = []
    for index, code in enumerate(top_codes, 1):
        row = by_code[code]
        top_rows.append((str(index), row["codigo"], row["nombre"], row["metric"], row["source_short"], row["action"]))
    add_table(
        doc,
        ["#", "Código", "EcoGesto", "Métrica", "Fuente", "Siguiente paso"],
        top_rows,
        [520, 1150, 2780, 1050, 2200, 1660],
        7.9,
    )


def add_presentation_defense(doc):
    add_heading(doc, "7. Guion de defensa para la presentación", 1)
    add_para(
        doc,
        "La explicación recomendada es directa: se ha construido un MVP funcional que calcula, registra y comunica impacto de forma trazable, pero no sustituye una auditoría externa. Esa frase baja el riesgo y, a la vez, demuestra madurez técnica.",
    )
    for item in [
        "No vendemos una certificación científica cerrada; vendemos un sistema preparado para ser auditado y mejorado.",
        "Cada EcoGesto tiene unidad, periodo, fuente, nivel de confianza y estado editorial.",
        "El usuario ve estimaciones útiles, no promesas absolutas. Los impactos reales dependen de ubicación, hábitos y tecnología sustituida.",
        "El sistema separa CO2e, agua, plástico y euros para evitar una suma ambiental falsa.",
        "La matriz permite que una persona experta revise primero los EcoGestos de más impacto.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Riesgos y mitigación", 2)
    add_table(
        doc,
        ["Riesgo", "Mitigación en el MVP"],
        [
            ("Sobreprometer exactitud científica.", "Usar niveles de confianza y disclaimer visible."),
            ("Sumar unidades incompatibles.", "Agregación separada por CO2e, agua, plástico y euros."),
            ("Factores desactualizados.", "Fecha de revisión y fuente guardada por resultado."),
            ("Alta variabilidad en alimentación/consumo.", "Rangos, supuestos explícitos y prioridad de revisión."),
            ("Impactos financieros indirectos.", "Mostrarlos como cualitativos o no agregables hasta metodología específica."),
        ],
        [3400, 5960],
        9.2,
        SOFT_GREEN,
    )


def add_resources(doc):
    add_heading(doc, "8. Recursos y fuentes consultadas", 1)
    add_para(
        doc,
        "Estas fuentes son la base recomendada para defender la metodología y completar la validación profesional. La columna de uso indica dónde encaja cada una dentro del MVP.",
    )
    sources = [
        (
            "MITECO - Calculadoras de huella de carbono",
            "https://www.miteco.gob.es/es/cambio-climatico/temas/mitigacion-politicas-y-medidas/calculadoras.html",
            "Marco español de factores y cálculo de huella, especialmente organización y emisiones energéticas/directas.",
        ),
        (
            "MITECO - Guía para el cálculo de la huella de carbono",
            "https://www.miteco.gob.es/content/dam/miteco/es/cambio-climatico/temas/mitigacion-politicas-y-medidas/guia_huella_carbono_tcm30-479093.pdf",
            "Base metodológica para explicar alcance, factores, límites y documentación de emisiones.",
        ),
        (
            "Red Eléctrica - Emisiones de CO2 del sistema eléctrico",
            "https://www.ree.es/es/datos/generacion/no-renovables-detalle-emisiones-CO2",
            "Factor eléctrico y lectura de kg/tCO2e por MWh para gestos de electricidad.",
        ),
        (
            "IDAE - Recomendaciones de ahorro energético en hogares",
            "https://www.idae.es/ahorra-energia/recomendaciones-para-el-ahorro-energetico-en-hogares",
            "Soporte para stand-by, iluminación, electrodomésticos, climatización y hábitos de ahorro.",
        ),
        (
            "IDAE - Guía práctica de la energía",
            "https://www.idae.es/uploads/documentos/12826.pdf",
            "Conducción eficiente, ahorro energético y recomendaciones domésticas.",
        ),
        (
            "ICAO - Carbon Emissions Calculator",
            "https://www.icao.int/environmental-protection/environmental-tools/icec",
            "Cálculo de emisiones de vuelos por trayecto y clase, mejor que un factor plano.",
        ),
        (
            "GHG Protocol - Scope 3 Calculation Guidance",
            "https://ghgprotocol.org/scope-3-calculation-guidance-2",
            "Marco para no simplificar en exceso cadenas de valor e impactos indirectos.",
        ),
        (
            "WRAP - Food waste",
            "https://www.wrap.ngo/what-we-do/future-proof-food/reducing-food-waste",
            "Soporte para desperdicio alimentario y peso climático de alimentos desperdiciados.",
        ),
        (
            "FAO - Tackling climate change through livestock",
            "https://www.fao.org/4/i3437e/i3437e.pdf",
            "Referencia de ganadería y emisiones de cadena alimentaria.",
        ),
        (
            "Agència Catalana de l'Aigua - Ahorro y eficiencia",
            "https://aca.gencat.cat/es/laca/campanyes-i-divulgacio/campanyes/estalvi-i-eficiencia-de-laigua/",
            "Benchmarks de consumo de agua doméstica y discurso de ahorro.",
        ),
        (
            "BOE - Ley 7/2022 de residuos y suelos contaminados",
            "https://www.boe.es/buscar/act.php?id=BOE-A-2022-5809",
            "Marco legal de prevención, reducción de residuos, reutilización y plásticos.",
        ),
    ]
    add_table(doc, ["Fuente", "URL", "Uso"], sources, [2250, 3400, 3710], 7.9, SOFT_GREEN)


def add_appendix_matrix(doc, rows):
    add_heading(doc, "Anexo A. Matriz de validación por EcoGesto", 1)
    add_para(
        doc,
        "La tabla resume la lectura metodológica de cada EcoGesto. No sustituye la revisión ambiental experta; sirve para priorizar y saber qué se puede defender ya como estimación y qué debe quedar pendiente.",
    )
    by_category = defaultdict(list)
    for row in rows:
        by_category[row["categoria"]].append(row)

    for category, category_rows in by_category.items():
        add_heading(doc, category, 2)
        table_rows = []
        for row in category_rows:
            table_rows.append(
                (
                    row["codigo"],
                    row["nombre"],
                    row["metric"],
                    row["confidence"],
                    row["source_family"],
                    row["action"],
                )
            )
        add_table(
            doc,
            ["Código", "EcoGesto", "Métrica", "Conf.", "Fuente", "Acción recomendada"],
            table_rows,
            [900, 2920, 900, 780, 1560, 2300],
            7.7,
            LIGHT_FILL,
        )


def add_appendix_disclaimer(doc):
    add_heading(doc, "Anexo B. Disclaimer recomendado", 1)
    add_callout(
        doc,
        "Texto para informes de usuario",
        "Los impactos mostrados son estimaciones orientativas basadas en fuentes y supuestos indicados en cada EcoGesto. Los resultados reales dependen del contexto, frecuencia, ubicación, tecnología sustituida y ejecución de la persona usuaria. Los indicadores no CO2 se muestran por separado y no deben interpretarse como equivalentes directos de CO2.",
        "F8FBF7",
    )
    add_para(
        doc,
        "Para el equipo AHC, la siguiente mejora natural es sustituir progresivamente cada rango de impacto por fichas de factor completas: variable de usuario, fórmula visible, factor, unidad, periodo, fuente, confianza, fecha de revisión y responsable.",
    )


def build_document():
    rows = load_data()
    category_counts = Counter(row["categoria"] for row in rows)
    high_priority = sum(1 for row in rows if row["prioridad_validacion"] == "Alta")

    doc = Document()
    configure_document(doc)

    add_cover(doc, len(rows), len(category_counts), high_priority, 10)
    add_executive_summary(doc)
    add_scope_and_architecture(doc)
    add_sources_method(doc)
    add_category_methodology(doc)
    add_validation_summary(doc, rows)
    add_presentation_defense(doc)
    add_resources(doc)
    add_appendix_matrix(doc, rows)
    add_appendix_disclaimer(doc)

    core = doc.core_properties
    core.title = "Informe de validación metodológica de cálculos de impacto - EcoGestos AHC"
    core.subject = "Lógica científica, fuentes, límites y matriz de defensa del MVP"
    core.author = "Codex para EcoGestos AHC"
    core.comments = "Generado a partir de la matriz de validación, ecogestos.json y la lógica de cálculo del MVP."
    core.created = datetime(2026, 6, 18, 12, 0, 0)
    core.modified = datetime(2026, 6, 18, 12, 0, 0)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_document()
    print(path)
